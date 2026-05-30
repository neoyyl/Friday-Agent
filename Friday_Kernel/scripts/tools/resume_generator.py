"""
生成个人简历（WPS/DOCX 格式）
目标：DeepSeek Agent 研发岗
—— 修复：移除所有可能导致方框的字符（emoji、特殊符号）
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 页面设置 ──
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ── 默认样式 ──
style = doc.styles["Normal"]
font = style.font
font.name = "微软雅黑"
font.size = Pt(10.5)
# 同时设置中文字体
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def set_font(run, name="微软雅黑", size=Pt(10.5), bold=False, color=None):
    """统一设置字体，同时指定 Latin 和 East-Asian 字体"""
    run.font.name = name
    run.font.size = size
    run.bold = bold
    rPr = run.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = run.element.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = color


def add_section_title(text):
    """添加模块标题（无特殊符号）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size=Pt(14), bold=True, color=RGBColor(0x1A, 0x1A, 0x2E))
    # 分隔线：用简单下划线替代特殊符号
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(6)
    run2 = p2.add_run("_" * 60)
    set_font(run2, size=Pt(6), color=RGBColor(0xBB, 0xBB, 0xBB))


def add_bullet_line(text, indent=Cm(0.3)):
    """添加带格式的项目行，用 - 代替 •"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = indent
    run = p.add_run("- " + text)
    set_font(run, size=Pt(10))


def add_project_block(title, subtitle, desc_lines):
    """项目经历条目"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run_t = p.add_run(title)
    set_font(run_t, size=Pt(11), bold=True)
    if subtitle:
        run_s = p.add_run("   |   " + subtitle)
        set_font(run_s, size=Pt(10), color=RGBColor(0x66, 0x66, 0x66))
    for line in desc_lines:
        add_bullet_line(line)


# ══════════════════════════════════════════
# 顶部：姓名 + 求职意向
# ══════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run = p.add_run("杨云龙")
set_font(run, size=Pt(24), bold=True, color=RGBColor(0x1A, 0x1A, 0x2E))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
run = p.add_run("求职意向：Agent 研发岗  |  期望城市：杭州")
set_font(run, size=Pt(12), color=RGBColor(0x33, 0x33, 0x33))

# ══════════════════════════════════════════
# 联系方式（纯文本，无emoji）
# ══════════════════════════════════════════
tbl = doc.add_table(rows=1, cols=1)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Table Grid"
cell = tbl.cell(0, 0)
cell.text = ""
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("电话: 18685566913    邮箱: 3182299867@qq.com    GitHub: github.com/yangyunlong")
set_font(run, size=Pt(10))

doc.add_paragraph()  # 空行

# ══════════════════════════════════════════
# 教育背景
# ══════════════════════════════════════════
add_section_title("教育背景")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(2)
run1 = p.add_run("石河子大学")
set_font(run1, size=Pt(11), bold=True)
run2 = p.add_run("    动物科学 本科    2022 - 2026")
set_font(run2, size=Pt(10.5))

# ══════════════════════════════════════════
# 核心项目
# ══════════════════════════════════════════
add_section_title("核心项目")

add_project_block(
    "Friday AI -- 全栈个人智能助手系统",
    "独立开发 / 核心架构设计",
    [
        '从零构建完整的 AI Agent 系统, 涵盖语音唤醒("Hey Friday"), 语音识别/TTS, 声纹识别, 多模态交互等全链路能力',
        "设计并实现 20 个子 Agent 智能调度架构(Master Agent > Judge > Dispatch > Execute > Audit), 支持链式/并行/条件路由等多种协作模式",
        "研发 GFCR(Generate-Filter-Control-Replay)记忆增强框架, 实现从会话中自动提取可复用行为模式, 让 Agent 具备持续自我进化能力",
        "开发 E3 推理增强框架(快速响应/分析/研究/决策/创意 5 模式), 接入 Chain-of-Thought, ReAct, Tree-of-Thought 等高级推理路径",
        "构建分发日志系统(Dispatch Logger), 自动统计调度成功率并给出优化建议, 形成数据驱动的迭代闭环",
        "集成系统监控(GPU/CPU/内存), 知识管理(Obsidian 双链笔记), 定时任务调度(6 个自动化任务)等功能模块",
        "发布 Friday-Portable v1.6.0, 打包为可独立分发的桌面版本",
    ],
)

add_project_block(
    "OpenCode 开源贡献",
    "GitHub / Agent 框架层",
    [
        "为 OpenCode/OpenClaw 开源 AI Agent 框架贡献代码, 涉及 Agent 技能调度与配置体系",
        "深度使用 OpenClaw 生态, 熟悉 Agent 技能(Skill)机制, 工具调用, 多 Agent 协作等核心模式",
    ],
)

# ══════════════════════════════════════════
# 技术能力
# ══════════════════════════════════════════
add_section_title("技术能力")

skills = [
    "编程语言: Python(主力), Shell, JavaScript/Node.js",
    "AI/Agent: LLM 应用开发, Agent 架构设计, RAG, Prompt Engineering, 多 Agent 协作, 工具调用(Tool Use)",
    "框架工具: OpenClaw/OpenCode 生态, LangChain 模式, python-docx, Flask, PyTorch(基础)",
    "系统能力: Windows/Linux 双平台, 进程管理, 系统监控, 桌面 GUI(Tkinter/PyQt), Git 版本控制",
    "硬件监控: GPU(NVIDIA), CPU, 内存, 磁盘实时监控与告警体系",
]
for s in skills:
    add_bullet_line(s)

# ══════════════════════════════════════════
# 开源贡献
# ══════════════════════════════════════════
add_section_title("开源贡献与社区")

items = [
    "OpenCode(OpenClaw)开源贡献: Agent 技能调度与配置方向代码提交",
    "深度用户: Claude Code, OpenClaw, Cursor 等 Agent 工具的重度使用者, 对 Agent 产品有第一手实践经验",
    "自行搭建并维护完整的 AI Agent 生产环境(含推理, 记忆, 调度, 监控全套链路)",
]
for item in items:
    add_bullet_line(item)

# ══════════════════════════════════════════
# 自我评价
# ══════════════════════════════════════════
add_section_title("自我评价")

self_eval = (
    "极具方向感与创造力, 能从零构建完整系统并持续迭代. 学习能力与研究能力强, "
    "能独立追踪前沿技术(如 Agent 架构, 推理增强, 记忆机制等)并快速落地为可工作的产品. "
    "热爱 AI Agent 方向, 对 OpenClaw 式个人智能助理, 多 Agent 协作, 记忆连续性等场景有深入实践. "
    "虽然是动物科学本科背景, 但这种跨学科视角带来了独特的系统思维与问题解决角度. "
    "期望在 DeepSeek 的 Agent 研发团队中, 与优秀的工程师和研究员一起, "
    "共同推进 AGI 时代下智能助理的产品化落地."
)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.first_line_indent = Cm(0.7)
run = p.add_run(self_eval)
set_font(run, size=Pt(10.5))

# ── 保存 ──
output_dir = os.path.expanduser("~/Desktop")
output_path = os.path.join(output_dir, "杨云龙-简历-DeepSeek-Agent研发岗.docx")
doc.save(output_path)
print(f"[OK] 简历已生成: {output_path}")
